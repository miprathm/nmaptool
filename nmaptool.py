from libnmap.process import NmapProcess
from libnmap.parser import NmapParser, NmapParserException
import sys, os
import re
import openpyxl
from docx import Document
from docx.shared import Inches


# start a new nmap scan on localhost with some specific options
def do_scan(targets, options):
    parsed = None
    nmproc = NmapProcess(targets, options)
    rc = nmproc.run()
    if rc != 0:
        print("nmap scan failed: {0}".format(nmproc.stderr))
    print(type(nmproc.stdout))
    file = open('current_ips.nmap','w')
    file.write(nmproc.stdout)
    file.close()
    try:
        parsed = NmapParser.parse(nmproc.stdout)
    except NmapParserException as e:
        print("Exception raised while parsing scan: {0}".format(e.msg))

    return parsed


# print scan results from a nmap report
def print_scan(nmap_report):
    print("Starting Nmap {0} ( http://nmap.org ) at {1}".format( nmap_report.version, nmap_report.started ))
    document = Document()
    document.add_heading('Nmap Scan', 0)	

    no_of_hosts = len(nmap_report.hosts)

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ips'
    hdr_cells[1].text = 'Ports'
    hdr_cells[2].text = 'Services'

    for host in nmap_report.hosts:
        row_cells = table.add_row().cells
        #ip
        row_cells[0].text = host.address
        string_port = ''
        string_service = ''
        for serv in host.services:
            string_port += str(serv.port) + "\n"
            string_service += serv.service + "\n"
        #Ports
        row_cells[1].text = string_port
        #Services
        row_cells[2].text = string_service
        """if len(host.hostnames):
            tmp_host = host.hostnames.pop()
        else:
            tmp_host = host.address

        print("Nmap scan report for {0} ({1})".format( tmp_host, host.address))
        print("Host is {0}.".format(host.status))
        print("  PORT     STATE         SERVICE")

        for serv in host.services:
            pserv = "{0:>5s}/{1:3s}  {2:12s}  {3}".format( str(serv.port),serv.protocol,serv.state,serv.service)
            if len(serv.banner):
                pserv += " ({0})".format(serv.banner)
            print(pserv)
        """   
    document.add_page_break()
    document.save('output.docx')
    print(nmap_report.summary)


if __name__ == "__main__":
    # Get Path from Cmd line
    pathname = os.path.abspath(sys.argv[0])
    if len(sys.argv) > 1 :
        pathname = os.path.abspath(sys.argv[1])
    print(pathname)

    all_ips_file = open(pathname,'r')

    all_ips = all_ips_file.read()

    ip_finder = re.compile(''' 
    	(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})\s*
    ''',re.X)

    ips = ip_finder.findall(all_ips)
    report = do_scan(ips, "-Pn -T4 -v")
    
    #report = do_scan("127.0.0.1", "-sV")
    if report :
        print_scan(report)
    else:
        print("No results returned")
